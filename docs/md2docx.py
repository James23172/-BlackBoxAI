#!/usr/bin/env python3
"""
P3 文档转换工具：Markdown → Word (.docx)
Mermaid 图表渲染为 PNG 嵌入文档

用法：python md2docx.py
"""

import re
import os
import sys
import base64
import zlib
import urllib.request
import urllib.error
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = None


def encode_mermaid_ink(code: str) -> str:
    """
    Mermaid.ink 编码: raw deflate + base64url
    参考 https://github.com/jihchi/mermaid.ink
    """
    code_bytes = code.encode('utf-8')
    compressor = zlib.compressobj(9, zlib.DEFLATED, -zlib.MAX_WBITS)
    compressed = compressor.compress(code_bytes) + compressor.flush()
    return base64.urlsafe_b64encode(compressed).rstrip(b'=').decode('ascii')


def render_mermaid_to_png(code: str, output_path: str) -> bool:
    """通过 mermaid.ink 渲染 Mermaid → PNG"""
    import time
    
    # 准备两种编码
    encode_methods = [
        lambda c: "https://mermaid.ink/img/" + encode_mermaid_ink(c),
        lambda c: "https://mermaid.ink/img/" + base64.urlsafe_b64encode(
            c.encode('utf-8')).rstrip(b'=').decode('ascii'),
    ]
    
    for attempt, url_fn in enumerate(encode_methods):
        try:
            url = url_fn(code)
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = resp.read()
                if len(data) > 200 and data[:4] == b'\x89PNG':
                    with open(output_path, 'wb') as f:
                        f.write(data)
                    return True
            if attempt == 0:
                time.sleep(0.3)
        except Exception:
            if attempt == 0:
                continue
    return False


def extract_mermaid_blocks(text: str):
    """提取 ```mermaid ... ``` 块"""
    pattern = re.compile(r'```mermaid\s*\n(.*?)```', re.DOTALL)
    blocks = {}
    bid = 0
    def repl(m):
        nonlocal bid
        blocks[bid] = m.group(1).strip()
        ph = f'[MERMAID_BLOCK_{bid}]'
        bid += 1
        return ph
    return pattern.sub(repl, text), blocks


def fix_chinese_font(run):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_code_block(doc, code_text: str):
    for line in code_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(16)
        pPr = p._element.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        run.font.size = Pt(9)
        run.font.name = 'Consolas'
        fix_chinese_font(run)


def add_heading_styled(doc, level, text):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        fix_chinese_font(run)


def add_image_centered(doc, image_path: str, max_w=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(max_w))


def convert_md_to_docx(input_path: str, output_path: str):
    print(f"\n[转换] {os.path.basename(input_path)}")
    
    with open(input_path, "r", encoding="utf-8") as f:
        content = f.read()
    
    img_dir = os.path.join(OUTPUT_DIR, "images")
    os.makedirs(img_dir, exist_ok=True)
    
    processed, mermaid_blocks = extract_mermaid_blocks(content)
    print(f"  发现 {len(mermaid_blocks)} 个 Mermaid 图表")
    
    # 渲染 Mermaid
    img_paths = {}
    for bid, code in mermaid_blocks.items():
        img_path = os.path.join(img_dir, f"mermaid_{bid}.png")
        print(f"  渲染图表 {bid+1}/{len(mermaid_blocks)}...")
        if render_mermaid_to_png(code, img_path):
            img_paths[bid] = img_path
        else:
            print(f"  [WARN] 图表 {bid+1} 渲染失败")
    
    # 创建 DOCX
    doc = Document()
    doc.styles['Normal'].font.size = Pt(11)
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)
    
    lines = processed.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []
    
    while i < len(lines):
        line = lines[i]
        
        # Mermaid placeholder
        mm = re.match(r'\[MERMAID_BLOCK_(\d+)\]', line)
        if mm:
            bid = int(mm.group(1))
            if bid in img_paths:
                add_image_centered(doc, img_paths[bid])
            else:
                # Render as Mermaid code block
                code = mermaid_blocks.get(bid, '')
                if code:
                    p = doc.add_paragraph()
                    run = p.add_run('[Mermaid 图表 - 请在支持 Mermaid 的工具中查看]')
                    run.font.size = Pt(9)
                    run.italic = True
                    fix_chinese_font(run)
                    add_code_block(doc, code)
            i += 1
            continue
        
        # Code block
        if line.startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Table
        if line.startswith('|') and not re.match(r'^\|[\s\-:]+\|', line.strip()):
            if not in_table:
                in_table = True
                table_buffer = []
            cells = [c.strip() for c in line.split('|')[1:-1]]
            table_buffer.append(cells)
            i += 1
            continue
        elif in_table:
            if table_buffer and len(table_buffer) >= 1:
                data_rows = table_buffer
                if len(data_rows) >= 2 and all(
                    c.replace('-','').replace(':','').strip() == '' 
                    for c in data_rows[0]
                ):
                    data_rows = data_rows[1:]
                num_cols = max(len(r) for r in data_rows) if data_rows else 1
                tbl = doc.add_table(rows=len(data_rows), cols=num_cols, style='Light Grid Accent 1')
                for ri, rd in enumerate(data_rows):
                    for ci, ct in enumerate(rd):
                        if ci < num_cols:
                            tbl.cell(ri, ci).text = ct
                            for para in tbl.cell(ri, ci).paragraphs:
                                for run in para.runs:
                                    run.font.size = Pt(9)
                doc.add_paragraph()
            table_buffer = []
            in_table = False
            continue
        
        # Heading
        hm = re.match(r'^(#{1,4})\s+(.+)$', line)
        if hm:
            add_heading_styled(doc, len(hm.group(1)), hm.group(2).strip())
            i += 1
            continue
        
        # Blockquote
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(line[2:])
            run.font.size = Pt(10)
            run.italic = True
            run.font.color.rgb = RGBColor(80, 80, 80)
            fix_chinese_font(run)
            i += 1
            continue
        
        # Unordered list
        ul = re.match(r'^(\s*)[-*+]\s+(.+)$', line)
        if ul:
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            run = p.add_run(ul.group(2))
            run.font.size = Pt(11)
            fix_chinese_font(run)
            i += 1
            continue
        
        # Ordered list
        ol = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ol:
            p = doc.add_paragraph(style='List Number')
            p.clear()
            run = p.add_run(ol.group(2))
            run.font.size = Pt(11)
            fix_chinese_font(run)
            i += 1
            continue
        
        # Horizontal rule
        if line.strip() in ('---', '***'):
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue
        
        # Empty line
        if not line.strip():
            i += 1
            continue
        
        # Normal paragraph
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = Pt(22)
        
        remaining = line
        while remaining:
            # Inline code
            cm = re.match(r'(.*?)`([^`]+)`(.*)', remaining)
            if cm:
                if cm.group(1):
                    run = p.add_run(cm.group(1))
                    run.font.size = Pt(11)
                    fix_chinese_font(run)
                run = p.add_run(cm.group(2))
                run.font.size = Pt(10)
                run.font.name = 'Consolas'
                fix_chinese_font(run)
                rPr = run._element.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'F0F0F0')
                rPr.append(shd)
                remaining = cm.group(3)
                continue
            
            # Bold
            bm = re.match(r'(.*?)\*\*(.+?)\*\*(.*)', remaining)
            if bm:
                if bm.group(1):
                    run = p.add_run(bm.group(1))
                    run.font.size = Pt(11)
                    fix_chinese_font(run)
                run = p.add_run(bm.group(2))
                run.font.size = Pt(11)
                run.bold = True
                fix_chinese_font(run)
                remaining = bm.group(3)
                continue
            
            run = p.add_run(remaining)
            run.font.size = Pt(11)
            fix_chinese_font(run)
            remaining = ''
        
        i += 1
    
    doc.save(output_path)
    print(f"  [OK] 已生成: {os.path.basename(output_path)}")


def main():
    global OUTPUT_DIR
    base_dir = r"C:\Users\wangwuguang\OneDrive\桌面\lab\1.0\docs"
    OUTPUT_DIR = base_dir
    
    files = [
        "P3-零基础完整教学文档.md",
        "P3-深度展开-完整技术剖析.md",
        "P3-代码完全注解版.md",
        "P3-图文精解-Mermaid图解版.md",
    ]
    
    print("使用 mermaid.ink 在线 API 渲染图表")
    
    for fname in files:
        fpath = os.path.join(base_dir, fname)
        if not os.path.exists(fpath):
            print(f"[WARN] 文件不存在: {fname}")
            continue
        output_path = os.path.join(base_dir, fname.replace(".md", ".docx"))
        try:
            convert_md_to_docx(fpath, output_path)
        except Exception as e:
            print(f"  [ERROR] 转换失败: {e}")
            import traceback
            traceback.print_exc()
    
    print("\n全部转换完成！")


if __name__ == "__main__":
    main()
